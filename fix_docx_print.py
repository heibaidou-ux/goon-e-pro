import docx
from docx.shared import Cm, Pt, Emu
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r'C:\Users\王晓东\xwechat_files\wxid_rxic3kxnrvkt12_67bf\msg\file\2026-05\2023年广州市高中阶段学校招生第三批次录取工分数线(1).docx'
OUT_PATH = r'C:\Users\王晓东\xwechat_files\wxid_rxic3kxnrvkt12_67bf\msg\file\2026-05\2023年广州市高中阶段学校招生第三批次录取工分数线(1)_A4打印版.docx'

d = docx.Document(DOC_PATH)

# ---- 1. Set A4 Landscape with tight margins ----
section = d.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.orientation = WD_ORIENT.LANDSCAPE

# Tight margins for maximum printable area
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

print(f"Page: A4 Landscape ({section.page_width.cm:.1f}x{section.page_height.cm:.1f}cm)")
print(f"Margins: L={section.left_margin.cm:.1f} R={section.right_margin.cm:.1f} T={section.top_margin.cm:.1f} B={section.bottom_margin.cm:.1f}cm")
print(f"Printable width: {(section.page_width - section.left_margin - section.right_margin).cm:.1f}cm")

# ---- 2. Process Table 1 (19 cols) ----
t = d.tables[0]
print(f"\nTable 1: {len(t.rows)} rows x {len(t.columns)} cols")

# Check if cols 15-19 (0-indexed) have data
col_has_data = {}
for j in range(14, 19):
    has_data = False
    for i in range(1, min(len(t.rows), 95)):
        if t.rows[i].cells[j].text.strip():
            has_data = True
            break
    col_has_data[j] = has_data
    print(f"  Col {j+1} has data: {has_data}")

# Reduce font size for entire table
for row in t.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(7.5)
                # Set Chinese font
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Set proportional column widths based on content type
# 19 cols: 序号(1) | 学校名称(2) | 户籍(3) | 范围(4) | 户籍生头(5-9) | 非户籍生头(10-14) | 外区生(15-19)
usable_width = section.page_width - section.left_margin - section.right_margin
available_width = usable_width.cm  # in cm

# Column widths in cm (proportional)
col_widths_cm = [
    0.8,    # 1: 序号
    3.0,    # 2: 学校名称 (wide - school names)
    0.8,    # 3: 户籍
    0.8,    # 4: 范围
    # 户籍生 5 col group
    1.1, 1.4, 1.1, 1.1, 1.4,
    # 非户籍生 5 col group
    1.1, 1.4, 1.1, 1.1, 1.4,
    # 外区生 5 col group (maybe empty)
    1.1, 1.4, 1.1, 1.1, 1.4,
]

# If cols 15-19 are empty, redistribute their width
total_planned = sum(col_widths_cm)
scale = available_width / total_planned
col_widths_cm = [w * scale for w in col_widths_cm]

# Set column widths
for j in range(len(t.columns)):
    for row in t.rows:
        row.cells[j].width = Cm(col_widths_cm[j])

print(f"\nAdjusted {len(t.columns)} columns to fit {available_width:.1f}cm width")
print(f"Column widths: {[f'{w:.1f}' for w in col_widths_cm]}")

# ---- 3. Process Table 2 (9 cols) - easier ----
t2 = d.tables[1]
print(f"\nTable 2: {len(t2.rows)} rows x {len(t2.columns)} cols")

# Reduce font size
for row in t2.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Column widths for 9 cols
col_widths_2 = [0.8, 4.5, 0.8, 0.8, 1.2, 1.5, 1.2, 1.2, 1.5]
total_2 = sum(col_widths_2)
scale_2 = available_width / total_2
col_widths_2 = [w * scale_2 for w in col_widths_2]

for j in range(len(t2.columns)):
    for row in t2.rows:
        row.cells[j].width = Cm(col_widths_2[j])

print(f"Adjusted {len(t2.columns)} columns")
print(f"Column widths: {[f'{w:.1f}' for w in col_widths_2]}")

# ---- 4. Save ----
d.save(OUT_PATH)
print(f"\nSaved to: {OUT_PATH}")
print("Done! Document is formatted for A4 landscape printing.")
